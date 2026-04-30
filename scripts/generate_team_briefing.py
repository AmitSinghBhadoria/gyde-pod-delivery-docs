"""
Generate Team Briefing Word document.
Professional internal doc for the engineering team.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# --- Color palette ---
DARK_BLUE = RGBColor(0x1B, 0x2A, 0x4A)
MEDIUM_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
ACCENT_GREEN = RGBColor(0x2D, 0x7D, 0x46)
ACCENT_ORANGE = RGBColor(0xE8, 0x6C, 0x00)
ACCENT_RED = RGBColor(0xCC, 0x33, 0x33)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            border = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{val["sz"]}" '
                f'w:space="0" w:color="{val["color"]}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)


def fmt(paragraph, text, bold=False, size=10, color=DARK_GRAY, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    return run


def add_heading(doc, text, level=1, color=DARK_BLUE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_styled_table(doc, headers, rows, header_color="1B2A4A", stripe=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt(p, h, bold=True, size=9, color=WHITE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.cell(r + 1, c)
            if stripe and r % 2 == 1:
                set_cell_shading(cell, "F5F7FA")
            p = cell.paragraphs[0]
            if isinstance(val, tuple):
                fmt(p, val[0], bold=True, size=9)
                fmt(p, val[1], size=9)
            else:
                fmt(p, str(val), size=9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    return table


def add_bullet(doc, text, bold_prefix=None, indent=Cm(1)):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        fmt(p, f"{bold_prefix}: ", bold=True, size=10)
        fmt(p, text, size=10)
    else:
        fmt(p, text, size=10)
    return p


def add_body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    fmt(p, text, size=10)
    return p


def add_callout_box(doc, title, text, bg_color="E8F0F8", border_color="2C5F8A"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg_color)
    set_cell_border(cell,
        left={"sz": "12", "color": border_color},
        top={"sz": "2", "color": border_color},
        bottom={"sz": "2", "color": border_color},
        right={"sz": "2", "color": border_color})
    p = cell.paragraphs[0]
    fmt(p, f"{title}  ", bold=True, size=10, color=MEDIUM_BLUE)
    fmt(p, text, size=9.5)


def build_doc():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ==========================================
    # TITLE PAGE
    # ==========================================

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "AI SUPPORT COPILOT", bold=True, size=28, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "Product Overview & Team Briefing", size=16, color=MEDIUM_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    fmt(p, "Phase 1 Pilot  |  Post-Discovery Update", size=12, color=MEDIUM_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    fmt(p, "Prepared by: Amit (POD Lead)", size=11, color=MEDIUM_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "Date: 2026-04-30", size=11, color=MEDIUM_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    fmt(p, "CONFIDENTIAL \u2014 Internal Team Only", bold=True, size=10, color=ACCENT_RED)

    doc.add_page_break()

    # ==========================================
    # TABLE OF CONTENTS (manual)
    # ==========================================

    add_heading(doc, "Contents", level=1)
    contents = [
        "1. The Problem We're Solving",
        "2. What We're Building",
        "3. Why AI (Not Traditional Software)",
        "4. Confirmed Technical Decisions",
        "5. The Dataset",
        "6. AI Pipeline Architecture",
        "7. Success Metrics & Targets",
        "8. Timeline & Milestones",
        "9. What Each Role Owns",
        "10. Key Constraints",
    ]
    for item in contents:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        fmt(p, item, size=11, color=MEDIUM_BLUE)

    doc.add_page_break()

    # ==========================================
    # 1. THE PROBLEM
    # ==========================================

    add_heading(doc, "1. The Problem We're Solving", level=1)

    add_body(doc, (
        "Our client's support team handles ~1,000 tickets per month through Freshdesk. "
        "When a ticket arrives, a human agent manually:"
    ))

    steps = [
        ("Reads and classifies", "Is this Authentication? Billing? How urgent?"),
        ("Searches the KB", "Opens the knowledge base, tries keywords, reads articles"),
        ("Decides what to do", "Reply with a fix? Ask for more info? Escalate to engineering?"),
        ("Checks escalation rules", "Which team? What context to include?"),
        ("Writes a response", "Drafts a reply that's accurate, cites the right article, matches the tone"),
    ]
    for prefix, text in steps:
        add_bullet(doc, text, bold_prefix=prefix)

    add_body(doc, "Every step is manual cognitive work. This causes:", space_after=4)

    add_styled_table(doc,
        ["Problem", "Impact"],
        [
            ["Slow", "Resolution time ranges from 2 hours to 2 days per ticket"],
            ["Inconsistent", "Different agents give different answers to the same question"],
            ["Error-prone", "Agents miss KB articles, apply wrong escalation rules"],
            ["Knowledge trapped", "Senior agents have judgment that junior agents lack"],
            ["Doesn't scale", "More tickets = more agents = linear cost growth"],
        ]
    )

    doc.add_paragraph()  # spacer

    add_callout_box(doc,
        "Core insight:",
        "The answers already exist in KB articles and escalation rules. "
        "The problem is finding and applying them consistently, every time, at speed. "
        "This is knowledge retrieval + judgment + writing \u2014 the exact combination where LLMs add value."
    )

    # ==========================================
    # 2. WHAT WE'RE BUILDING
    # ==========================================

    doc.add_page_break()
    add_heading(doc, "2. What We're Building", level=1)

    add_body(doc, (
        "An AI copilot that helps support agents resolve tickets faster. For the pilot, this is a "
        "standalone web application with an integrated copilot sidebar. Post-pilot, the same backend "
        "powers a Chrome extension / Freshdesk sidebar widget."
    ))

    add_heading(doc, "Pilot UI: Support Dashboard with Copilot Sidebar", level=2)

    add_body(doc, "The pilot is a three-panel web application:", space_after=4)

    add_styled_table(doc,
        ["Panel", "What It Shows", "Notes"],
        [
            ["Left: Ticket Queue", "All tickets from the dataset, filterable by category/priority/status",
             "Loaded from Excel for pilot; Freshworks API in production"],
            ["Center: Ticket Detail", "Full ticket view \u2014 subject, description, customer, channel, history",
             "Opens when agent clicks a ticket from the queue"],
            ["Right: Copilot Sidebar", "Classification, KB matches, recommended action, draft response, confidence",
             "Auto-triggers when a ticket is selected"],
        ]
    )

    doc.add_paragraph()

    add_heading(doc, "How an Agent Uses It (Step by Step)", level=2)

    agent_flow = [
        ("Opens the app", "Sees the ticket queue with all pending tickets"),
        ("Clicks a ticket", "Ticket detail opens in the center panel; copilot sidebar auto-processes"),
        ("Reviews classification", "Category, priority, sentiment \u2014 with confidence score"),
        ("Reads KB matches", "Top relevant articles with relevance scores \u2014 can expand to read full content"),
        ("Checks recommended action", "Reply / Ask for more info / Escalate \u2014 with reasoning explanation"),
        ("Reviews draft response", "AI-generated response grounded in KB articles with citations"),
        ("Accepts, edits, or overrides", "Agent always has final say \u2014 can modify the draft or write their own"),
        ("Rates the response", "\"Was this helpful?\" \u2014 feedback is stored for system improvement"),
    ]
    for prefix, text in agent_flow:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_paragraph()

    add_callout_box(doc,
        "Key principle:",
        "The copilot is NOT a chatbot. The agent does not type questions or have a conversation. "
        "The copilot automatically analyzes the ticket and presents its recommendations. "
        "The agent's job is to review, decide, and act.",
        bg_color="E8F0F8", border_color="2C5F8A"
    )

    doc.add_paragraph()

    add_heading(doc, "Post-Pilot: Chrome Extension on Freshdesk", level=2)

    add_body(doc, (
        "After the pilot, the same backend API powers a Chrome extension that overlays on Freshdesk. "
        "The agent opens a ticket in Freshdesk, the extension detects it, calls the API, and shows "
        "the copilot sidebar. No change to the copilot logic \u2014 only the trigger surface changes."
    ))

    add_heading(doc, "What the copilot does for every ticket", level=2)

    add_styled_table(doc,
        ["Step", "Copilot Does", "Agent Used To Do"],
        [
            ["Classify", "Reads ticket, assigns category + priority + sentiment", "Mentally categorize"],
            ["Retrieve", "Finds most relevant KB articles with citations and scores", "Manually search KB"],
            ["Recommend", "Suggests: Reply / Ask for info / Escalate \u2014 with reasoning", "Check escalation rules mentally"],
            ["Draft", "Writes grounded response citing KB, appropriate tone", "Write from scratch"],
            ["Present", "Shows sources, confidence score, reasoning trace", "Reasoning stays in agent's head"],
        ]
    )

    doc.add_paragraph()  # spacer

    add_heading(doc, "What the copilot does NOT do", level=2)

    not_items = [
        "Auto-send responses. Human-in-the-loop always \u2014 agent reviews, edits, decides",
        "Replace agents. The copilot assists, the agent is in control",
        "Handle cases outside the support domain",
        "Make autonomous decisions with legal or financial consequences",
    ]
    for item in not_items:
        add_bullet(doc, item)

    # ==========================================
    # 3. WHY AI
    # ==========================================

    add_heading(doc, "3. Why AI (Not Traditional Software)", level=1)

    add_body(doc, "Could we solve this with decision trees, keyword search, and templates? Partially. But it breaks because:")

    reasons = [
        ("Natural language is the input",
         "\"OTP not arriving\", \"MFA not working\", and \"can't log in, no SMS\" are the same issue. "
         "Keyword matching fails on this variation. LLMs understand semantic equivalence."),
        ("KB-to-ticket mapping is fuzzy",
         "A single ticket might match multiple KB articles. Escalation depends on combining ticket content + "
         "KB articles + escalation rules + context. This multi-factor reasoning is what LLMs handle well."),
        ("Responses aren't templatable",
         "Tone shifts with sentiment. Required info changes per escalation rule. Too many combinations for templates."),
        ("The system improves with data",
         "As KB articles are added or feedback is collected, the copilot adapts without code changes."),
    ]
    for prefix, text in reasons:
        add_bullet(doc, text, bold_prefix=prefix)

    # ==========================================
    # 4. CONFIRMED DECISIONS
    # ==========================================

    doc.add_page_break()
    add_heading(doc, "4. Confirmed Technical Decisions", level=1)

    add_body(doc, "The following decisions were confirmed during the Discovery call with Prasanna on April 29:")

    add_styled_table(doc,
        ["Decision", "Answer", "Notes"],
        [
            ["Cloud platform", "GCP", "Our own account for pilot"],
            ["LLM provider", "Our recommendation", "Must be LLM-agnostic \u2014 swappable"],
            ["Data source (pilot)", "Excel dataset", "36 tickets, 12 KB articles, 5 rules"],
            ["Data source (production)", "Freshworks APIs", "Architect for it, don't build yet"],
            ["Deployment (pilot)", "Standalone web app with copilot sidebar", "Three-panel dashboard"],
            ["Deployment (production)", "Chrome extension on Freshdesk", "Same API, different trigger surface"],
            ["Post-pilot handover", "Full on-premises", "Code, models, data \u2014 everything"],
            ["Communication", "1 weekly email + 1 call", "Email + Google Meet"],
            ["Accuracy target", "85%", "Classification, retrieval, action accuracy"],
            ["Auto-answer target", "30% of recurring tickets", "~300 tickets/month"],
            ["Go-live validation", "1000 synthetic questions", "We generate, their team lead reviews"],
        ]
    )

    # ==========================================
    # 5. THE DATASET
    # ==========================================

    doc.add_paragraph()
    add_heading(doc, "5. The Dataset", level=1)

    add_body(doc, "Prasanna shared a starter dataset with 4 data sheets:")

    add_styled_table(doc,
        ["Sheet", "Records", "Purpose", "Key Fields"],
        [
            ["Tickets_Historical", "36", "Past resolved tickets for building and testing",
             "ticket_id, subject, description, category, priority, sentiment, channel, source_kb_id, action_taken, resolution_summary"],
            ["KB_Articles", "12", "Knowledge base \u2014 what the copilot retrieves from",
             "kb_id, title, category, content, keywords, agent_notes"],
            ["Escalation_Rules", "5", "Rules mapping conditions to escalation teams",
             "rule_id, condition, escalation_team, required_context, sla_hours"],
            ["Evaluation_Set", "12", "Held-out blind test set \u2014 DO NOT use for training",
             "eval_id, ticket_text, expected_category, expected_priority, expected_action, expected_kb_id"],
        ]
    )

    doc.add_paragraph()

    add_heading(doc, "Coverage", level=2)

    add_styled_table(doc,
        ["Dimension", "Values"],
        [
            ["Categories (7)", "Authentication, Billing, Data Import, Integrations, Access Control, Compliance, Known Issue"],
            ["Priorities (4)", "Critical, High, Medium, Low"],
            ["Channels (3)", "Email, Chat, Portal"],
            ["Sentiments (3)", "Frustrated, Neutral, Satisfied"],
            ["Actions (3)", "Reply, Ask for more info, Escalate"],
            ["Escalation Teams (5)", "Engineering, Integrations Eng, Finance Ops, Compliance, Platform Ops"],
        ]
    )

    doc.add_paragraph()

    add_callout_box(doc,
        "Data quality note:",
        "The dataset has 11 unique ticket scenarios across 36 records (metadata varies but descriptions repeat). "
        "This is sufficient for POC architecture validation but will need synthetic expansion for meaningful evaluation. "
        "See the Data Feasibility Report for full analysis.",
        bg_color="FFF8E8", border_color="E86C00"
    )

    # ==========================================
    # 6. AI PIPELINE
    # ==========================================

    doc.add_page_break()
    add_heading(doc, "6. AI Pipeline Architecture", level=1)

    add_body(doc, "The copilot processes every ticket through a 5-step pipeline. Each step is independently testable and measurable.")

    pipeline_steps = [
        ["1. CLASSIFY", "LLM reads ticket text", "Category + Priority + Sentiment",
         "Prompt-based classification with structured output"],
        ["2. RETRIEVE", "Embed ticket \u2192 search KB index", "Top-K relevant KB articles with scores",
         "Hybrid retrieval: vector similarity + BM25 keyword matching"],
        ["3. REASON", "LLM receives ticket + KB + escalation rules", "Action: Reply / Ask for info / Escalate",
         "If escalate: which team + required context per rule"],
        ["4. DRAFT", "LLM generates grounded response", "Draft response with KB citations",
         "Tone-appropriate, includes confidence score"],
        ["5. PRESENT", "Display to agent in sidebar", "Full copilot output with reasoning",
         "Agent reviews, edits, sends (or overrides)"],
    ]

    add_styled_table(doc,
        ["Step", "What Happens", "Output", "Technical Notes"],
        pipeline_steps
    )

    doc.add_paragraph()

    add_heading(doc, "Feedback Loop (Client Requirement)", level=2)

    add_body(doc, (
        "After the agent processes a ticket with the copilot's help, the system captures feedback:"
    ))

    feedback_steps = [
        "Agent rates the response (\"was it helpful?\")",
        "Agent can edit the draft before sending",
        "The edited version is saved as a correct response",
        "The system learns from this feedback over time, creating golden evaluations from real usage",
    ]
    for i, step in enumerate(feedback_steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(2)
        fmt(p, f"{i}. ", bold=True, size=10, color=MEDIUM_BLUE)
        fmt(p, step, size=10)

    # ==========================================
    # 7. SUCCESS METRICS
    # ==========================================

    add_heading(doc, "7. Success Metrics & Targets", level=1)

    add_styled_table(doc,
        ["#", "Metric", "Type", "Target", "How Measured"],
        [
            ["1", "Classification accuracy", "AI Quality", "\u2265 85%", "Exact match on category + priority"],
            ["2", "Retrieval accuracy", "AI Quality", "\u2265 85%", "Expected KB article in retrieved set"],
            ["3", "Action accuracy", "AI Quality", "\u2265 85%", "Exact match on recommended action"],
            ["4", "Response acceptance rate", "Business", "\u2265 70%", "% drafts accepted without major edits"],
            ["5", "Auto-answer coverage", "Business", "30%", "% of recurring tickets handled autonomously"],
            ["6", "Autonomous operation", "Business", "No handholding", "System answers without manual intervention"],
        ]
    )

    doc.add_paragraph()

    add_body(doc, (
        "Go-live validation: Before going live, the copilot will be evaluated against 1,000 synthetic questions "
        "generated from the existing dataset. The client's support team lead will review the results."
    ))

    # ==========================================
    # 8. TIMELINE
    # ==========================================

    add_heading(doc, "8. Timeline & Milestones", level=1)

    add_callout_box(doc,
        "Hard deadlines:",
        "Sprint 1 Demo: May 10, 2026  |  Final Delivery: May 16, 2026. These are not flexible.",
        bg_color="FFF0F0", border_color="CC3333"
    )

    doc.add_paragraph()

    add_styled_table(doc,
        ["Sprint", "Dates", "Goal", "Key Deliverables"],
        [
            ["Sprint 1", "May 1 \u2013 May 10", "Walking skeleton + Eval harness",
             "End-to-end pipeline working, eval harness operational, baseline metrics, Sprint 1 demo"],
            ["Sprint 2", "May 11 \u2013 May 16", "MVP + Hardening + Handover",
             "Feedback loop, guardrails, 1000-question eval, UI polish, all documentation, final delivery"],
        ]
    )

    # ==========================================
    # 9. ROLE OWNERSHIP
    # ==========================================

    doc.add_page_break()
    add_heading(doc, "9. What Each Role Owns", level=1)

    add_body(doc, "Every team member has specific responsibilities in the pipeline and the engagement.")

    # Atharva
    add_heading(doc, "Atharva \u2014 AI Engineer", level=2)
    atharva_items = [
        "LLM prompts for classification, reasoning, and response drafting",
        "Retrieval pipeline: embedding model, hybrid search (vector + BM25), reranking",
        "LLM gateway: provider-agnostic abstraction so we can swap models without code changes",
        "Confidence scoring mechanism",
        "Feedback loop implementation (agent rates/edits, system stores corrections)",
        "Integration with Nancy's data pipeline (embeddings, vector store queries)",
    ]
    for item in atharva_items:
        add_bullet(doc, item)

    # Nancy
    add_heading(doc, "Nancy \u2014 Data Engineer", level=2)
    nancy_items = [
        "Data ingestion pipeline: Excel \u2192 structured data (pilot); Freshworks APIs (production architecture)",
        "KB article chunking, embedding generation, and vector store indexing",
        "Data quality checks on the dataset (see Data Feasibility Report for known issues)",
        "Schema design for ticket and KB storage",
        "Fix 3 null source_kb_id values in the dataset",
        "Support Nishka with golden dataset expansion",
    ]
    for item in nancy_items:
        add_bullet(doc, item)

    # Nishka
    add_heading(doc, "Nishka \u2014 QA", level=2)
    nishka_items = [
        "Evaluation harness: automated scoring pipeline against the golden dataset",
        "Golden dataset curation: expand from 12 held-out cases to 30-40 for Sprint 1",
        "Synthetic question generation: 1,000 questions for go-live validation",
        "Adversarial test cases: prompt injection, out-of-scope tickets, edge cases",
        "Regression testing: track metric changes across prompt/model updates",
        "End-to-end smoke tests for every ticket category",
    ]
    for item in nishka_items:
        add_bullet(doc, item)

    # Shubham
    add_heading(doc, "Shubham \u2014 Governance Engineer", level=2)
    shubham_items = [
        "Threat model: prompt injection, PII leakage, output manipulation, data exfiltration",
        "Guardrails: profanity filtering, misuse prevention, adversarial input handling",
        "Secrets management: LLM API keys, database credentials \u2014 none in code",
        "Data classification for production (not needed for pilot, but document requirements)",
        "Pre-release security review and sign-off before final delivery",
        "Responsible AI checklist: AI disclosure, confidence disclaimers, bias assessment",
    ]
    for item in shubham_items:
        add_bullet(doc, item)

    # Shivani
    add_heading(doc, "Shivani \u2014 Implementation Manager", level=2)
    shivani_items = [
        "POD Charter: draft, get Prasanna's sign-off before Sprint 1",
        "Sprint planning: backlog, story estimation, sprint goals",
        "Weekly status email to Prasanna + weekly Google Meet call",
        "Risk register: maintain and update throughout the engagement",
        "Scope management: any change requests go through Shivani",
        "Final delivery coordination: demos, documentation handover, walkthrough",
    ]
    for item in shivani_items:
        add_bullet(doc, item)

    # Amit
    add_heading(doc, "Amit \u2014 POD Lead", level=2)
    amit_items = [
        "Architecture decisions: system design, tech stack, ADRs",
        "Pilot web app UI: three-panel dashboard (ticket queue + detail + copilot sidebar)",
        "Code review: final say on all AI components",
        "Evaluation strategy: metric definitions, threshold proposals",
        "Technical demos and client-facing architecture discussions",
        "Knowledge transfer package and productionization note",
    ]
    for item in amit_items:
        add_bullet(doc, item)

    # ==========================================
    # 10. KEY CONSTRAINTS
    # ==========================================

    add_heading(doc, "10. Key Constraints", level=1)

    add_styled_table(doc,
        ["Constraint", "Detail", "Impact"],
        [
            ["Human-in-the-loop", "Copilot suggests, agent decides. No auto-send ever.", "UI must always allow edit/override"],
            ["LLM agnostic", "Must support swapping LLM providers", "Build provider abstraction layer"],
            ["On-premises handover", "Post-pilot: all code, models, data go to client", "No proprietary dependencies, document everything"],
            ["Profanity / misuse", "Zero tolerance on outputs", "Guardrails layer before every response"],
            ["Grounded responses", "Must cite KB sources, no hallucination", "Response generation always references retrieved articles"],
            ["Hard deadlines", "Demo: May 10, Delivery: May 16", "No scope creep \u2014 prioritize ruthlessly"],
            ["Documentation = deliverable", "Docs are as important as working code", "Document as you build, not after"],
        ]
    )

    doc.add_paragraph()  # spacer

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    fmt(p, "\u2014  End of Document  \u2014", size=10, color=MEDIUM_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "Gyde AI POD  |  Confidential  |  2026-04-30", size=9, color=MEDIUM_GRAY)

    return doc


if __name__ == "__main__":
    doc = build_doc()

    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "client-docs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Team_Briefing_AI_Support_Copilot.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")

    sim_dir = "/Users/amit/Work/Gyde/Gyde Pivot/Gyde AI POD Framework/Simulation Docs"
    sim_path = os.path.join(sim_dir, "Team_Briefing_AI_Support_Copilot.docx")
    doc.save(sim_path)
    print(f"Saved: {sim_path}")
