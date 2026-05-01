"""
Generate Architecture Sketch Word document.
Client-facing document with system design.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

DARK_BLUE = RGBColor(0x1B, 0x2A, 0x4A)
MEDIUM_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
ACCENT_GREEN = RGBColor(0x2D, 0x7D, 0x46)
ACCENT_ORANGE = RGBColor(0xE8, 0x6C, 0x00)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def fmt(paragraph, text, bold=False, size=10, color=DARK_GRAY, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    return run


def add_heading(doc, text, level=1, color=DARK_BLUE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_styled_table(doc, headers, rows, header_color="1B2A4A", stripe=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt(p, h, bold=True, size=9, color=WHITE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.cell(r + 1, c)
            if stripe and r % 2 == 1:
                set_cell_shading(cell, "F5F7FA")
            p = cell.paragraphs[0]
            fmt(p, str(val), size=9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    return table


def add_body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    fmt(p, text, size=10)
    return p


def add_bullet(doc, text, bold_prefix=None, indent=Cm(1)):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if bold_prefix:
        fmt(p, f"{bold_prefix}: ", bold=True, size=9.5)
        fmt(p, text, size=9.5)
    else:
        fmt(p, text, size=9.5)
    return p


def add_callout(doc, title, text, bg="E8F0F8", border="2C5F8A"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="{border}"/>'
        f'<w:top w:val="single" w:sz="2" w:space="0" w:color="{border}"/>'
        f'<w:bottom w:val="single" w:sz="2" w:space="0" w:color="{border}"/>'
        f'<w:right w:val="single" w:sz="2" w:space="0" w:color="{border}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    p = cell.paragraphs[0]
    fmt(p, f"{title}  ", bold=True, size=10, color=MEDIUM_BLUE)
    fmt(p, text, size=9.5)


def build_doc():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ===== TITLE =====
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "ARCHITECTURE SKETCH", bold=True, size=28, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "AI Support Copilot \u2014 Phase 1 Pilot", size=16, color=MEDIUM_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    fmt(p, "Version 1.0  |  2026-05-01", size=11, color=MEDIUM_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    fmt(p, "Owner: Amit (POD Lead)", size=11, color=MEDIUM_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    fmt(p, "CONFIDENTIAL", bold=True, size=10, color=RGBColor(0xCC, 0x33, 0x33))

    doc.add_page_break()

    # ===== SYSTEM OVERVIEW =====
    add_heading(doc, "1. System Overview", level=1)

    add_body(doc, (
        "The AI Support Copilot is a RAG-based (Retrieval-Augmented Generation) system with a "
        "sequential pipeline architecture. It receives a support ticket, classifies it, retrieves "
        "relevant KB articles via hybrid search, reasons over the combined context to recommend "
        "an action, and drafts a grounded response \u2014 all presented to a human agent for review."
    ))

    add_callout(doc, "Architecture Pattern: RAG",
        "Grounded responses (no hallucination), no training data needed, KB updates without "
        "retraining, every response cites sources, fully portable for on-prem handover.")

    # ===== ARCHITECTURE DIAGRAM =====
    doc.add_paragraph()
    add_heading(doc, "2. System Architecture", level=1)

    arch_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "client-docs", "architecture_diagram.png")
    if os.path.exists(arch_path):
        doc.add_picture(arch_path, width=Inches(6.2))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt(p, "Figure 1: System Architecture \u2014 AI pipeline, data flow, and infrastructure",
            size=8, color=MEDIUM_GRAY)

    # ===== TECH STACK =====
    doc.add_paragraph()
    add_heading(doc, "3. Technology Stack", level=1)

    add_styled_table(doc,
        ["Layer", "Technology", "Rationale"],
        [
            ["Cloud", "GCP (Google Service Account)", "Client's cloud; single identity for all services"],
            ["LLM", "Gemini via Vertex AI", "Native GCP integration; strong structured output"],
            ["Embeddings", "Vertex AI text-embedding-005 (768d)", "Same service account; no extra provider"],
            ["Backend", "Express (Node.js)", "Team expertise; fast iteration"],
            ["Frontend", "React", "Component model fits three-panel layout"],
            ["Orchestration", "LangChain.js", "Provider abstraction; structured output; ES + MongoDB support"],
            ["Operational DB", "MongoDB", "Flexible schema; tickets, feedback, audit logs"],
            ["Search / Retrieval", "Elasticsearch (hybrid)", "Vector + BM25 in single query; self-hostable"],
            ["Search Fusion", "Reciprocal Rank Fusion (RRF)", "Balanced merge of vector + keyword; no tuning"],
        ]
    )

    # ===== AI PIPELINE =====
    doc.add_page_break()
    add_heading(doc, "4. AI Pipeline (LangChain.js)", level=1)

    add_body(doc, "Five sequential steps, each independently testable and measurable:")

    add_styled_table(doc,
        ["Step", "Model / Engine", "Input", "Output", "Technique"],
        [
            ["1. Classify", "Gemini (Vertex AI)",
             "Ticket text",
             "Category + Priority + Sentiment + Confidence",
             "Structured output (JSON schema)"],
            ["2. Retrieve", "Elasticsearch",
             "Ticket text (embedded)",
             "Top-K KB articles + relevance scores",
             "Hybrid: kNN vector + BM25, merged via RRF"],
            ["3. Reason", "Gemini (Vertex AI)",
             "Ticket + KB articles + Escalation rules",
             "Action (Reply/Ask/Escalate) + Reasoning + Confidence",
             "Chain-of-thought prompting"],
            ["4. Draft", "Gemini (Vertex AI)",
             "Ticket + KB + Action + Reasoning",
             "Grounded response with KB citations",
             "RAG with citation tracking"],
            ["5. Guardrails", "Post-processing",
             "All pipeline outputs",
             "Guardrail status + warnings",
             "Profanity filter, PII check, confidence gating"],
        ]
    )

    # ===== DATA LAYER =====
    doc.add_paragraph()
    add_heading(doc, "5. Data Layer", level=1)

    add_heading(doc, "MongoDB Collections", level=2)
    add_styled_table(doc,
        ["Collection", "Purpose", "Key Fields"],
        [
            ["tickets", "Ingested ticket data", "ticket_id, subject, description, category, priority, channel"],
            ["feedback", "Agent corrections and ratings", "ticket_id, helpful, original_draft, edited_draft"],
            ["audit_log", "Complete copilot decision trail", "ticket_id, pipeline_output, latency_ms, model_version"],
            ["sessions", "Agent session tracking", "agent_id, session_start, tickets_processed"],
        ]
    )

    doc.add_paragraph()
    add_heading(doc, "Elasticsearch Indices", level=2)
    add_styled_table(doc,
        ["Index", "Purpose", "Configuration"],
        [
            ["kb_articles", "KB storage + BM25 search", "Standard analyzer; fields: kb_id, title, content, keywords"],
            ["kb_vectors", "KB embeddings for kNN search", "HNSW index, cosine similarity, 768 dimensions"],
        ]
    )

    # ===== LLM GATEWAY =====
    doc.add_paragraph()
    add_heading(doc, "6. LLM Gateway (Provider-Agnostic)", level=1)

    add_body(doc, (
        "The LLM Gateway provides a uniform interface for all LLM calls. The active provider "
        "is selected via environment variable (LLM_PROVIDER). Swapping from Gemini to GPT-4o "
        "or Claude requires zero code changes \u2014 only a config update."
    ))

    add_styled_table(doc,
        ["Provider", "Status", "Use Case"],
        [
            ["Vertex AI (Gemini)", "Primary (default)", "All pipeline steps; native GCP integration"],
            ["OpenAI (GPT-4o)", "Fallback", "If Gemini accuracy is insufficient on specific steps"],
            ["Anthropic (Claude)", "Fallback", "Alternative fallback; strong reasoning capabilities"],
        ]
    )

    # ===== INTEGRATION POINTS =====
    doc.add_paragraph()
    add_heading(doc, "7. Integration Points", level=1)

    add_styled_table(doc,
        ["Integration", "Pilot", "Production"],
        [
            ["Ticket source", "Excel \u2192 MongoDB (ingestion script)", "Freshworks API \u2192 MongoDB (webhook/polling)"],
            ["KB source", "Excel \u2192 Elasticsearch (indexing script)", "Freshworks KB API \u2192 Elasticsearch (scheduled)"],
            ["Escalation rules", "JSON config loaded at startup", "Freshworks or config service"],
            ["LLM", "Vertex AI (Gemini) via Service Account", "Same; or swap via LLM Gateway config"],
            ["Agent interface", "Standalone React web app", "Chrome extension calling same backend API"],
        ]
    )

    # ===== NON-FUNCTIONAL =====
    doc.add_paragraph()
    add_heading(doc, "8. Non-Functional Requirements", level=1)

    add_styled_table(doc,
        ["Requirement", "Target", "Approach"],
        [
            ["Latency", "< 10s per ticket (full pipeline)", "Parallelize classify + retrieve; streaming where possible"],
            ["Portability", "Full on-prem handover", "All components self-hostable; no managed-only services"],
            ["LLM agnostic", "Swap provider via config", "LLM Gateway abstraction; env var selection"],
            ["Auditability", "Every decision traceable", "Audit log with full pipeline output per ticket"],
            ["Security", "No secrets in code", "GCP Service Account; env vars for config"],
            ["Versioning", "Prompts + data in Git", "Prompt templates as files; dataset versioned in repo"],
        ]
    )

    # ===== GCP SETUP =====
    doc.add_paragraph()
    add_heading(doc, "9. GCP Service Account Setup", level=1)

    add_styled_table(doc,
        ["Permission / Role", "Service", "Purpose"],
        [
            ["roles/aiplatform.user", "Vertex AI", "Gemini LLM + embedding API calls"],
            ["roles/serviceusage.serviceUsageConsumer", "Service Usage", "Enable required APIs"],
            ["Enable: aiplatform.googleapis.com", "Vertex AI API", "Required for Gemini and embeddings"],
        ]
    )

    doc.add_paragraph()
    add_body(doc, (
        "MongoDB and Elasticsearch run on the existing GCP VM. No additional IAM roles needed. "
        "Firewall rules restrict access to internal traffic only."
    ))

    # ===== WALKING SKELETON =====
    doc.add_paragraph()
    add_heading(doc, "10. Walking Skeleton (Sprint 1 Target)", level=1)

    add_callout(doc, "Thinnest end-to-end slice:",
        "One hardcoded ticket \u2192 Gemini classifies \u2192 Elasticsearch retrieves top-1 KB "
        "\u2192 Gemini reasons + drafts \u2192 React sidebar displays full output. "
        "When this works, the architecture is validated. Everything after is iteration.")

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "\u2014  End of Document  \u2014", size=10, color=MEDIUM_GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "Gyde AI POD  |  Confidential  |  2026-05-01", size=9, color=MEDIUM_GRAY)

    return doc


if __name__ == "__main__":
    doc = build_doc()

    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "client-docs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Architecture_Sketch_AI_Support_Copilot.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")

    sim_dir = "/Users/amit/Work/Gyde/Gyde Pivot/Gyde AI POD Framework/Simulation Docs"
    sim_path = os.path.join(sim_dir, "Architecture_Sketch_AI_Support_Copilot.docx")
    doc.save(sim_path)
    print(f"Saved: {sim_path}")
