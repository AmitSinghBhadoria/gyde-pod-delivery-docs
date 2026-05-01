"""
Generate POD Charter Word document.
Client-facing document for sign-off.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

DARK_BLUE = RGBColor(0x1B, 0x2A, 0x4A)
MEDIUM_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
ACCENT_GREEN = RGBColor(0x2D, 0x7D, 0x46)
ACCENT_ORANGE = RGBColor(0xE8, 0x6C, 0x00)
ACCENT_RED = RGBColor(0xCC, 0x33, 0x33)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def fmt(paragraph, text, bold=False, size=10, color=DARK_GRAY, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    return run


def add_heading(doc, text, level=1, color=DARK_BLUE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_styled_table(doc, headers, rows, header_color="1B2A4A", stripe=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt(p, h, bold=True, size=9, color=WHITE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.cell(r + 1, c)
            if stripe and r % 2 == 1:
                set_cell_shading(cell, "F5F7FA")
            p = cell.paragraphs[0]
            fmt(p, str(val), size=9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    return table


def add_body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    fmt(p, text, size=10)
    return p


def add_bullet(doc, text, bold_prefix=None, indent=Cm(1)):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if bold_prefix:
        fmt(p, f"{bold_prefix}: ", bold=True, size=9.5)
        fmt(p, text, size=9.5)
    else:
        fmt(p, text, size=9.5)
    return p


def add_callout(doc, text, bg="E8F0F8", border="2C5F8A"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="{border}"/>'
        f'<w:top w:val="single" w:sz="2" w:space="0" w:color="{border}"/>'
        f'<w:bottom w:val="single" w:sz="2" w:space="0" w:color="{border}"/>'
        f'<w:right w:val="single" w:sz="2" w:space="0" w:color="{border}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    p = cell.paragraphs[0]
    fmt(p, text, size=9.5)


def build_doc():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ===== TITLE =====
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "POD CHARTER", bold=True, size=28, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "AI Support Copilot \u2014 Phase 1 Pilot", size=16, color=MEDIUM_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    fmt(p, "Version 1.0  |  2026-04-30", size=11, color=MEDIUM_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    fmt(p, "Gyde AI POD Framework  |  Doc 01, Section 6", size=10, color=MEDIUM_GRAY)

    # Signatories box
    doc.add_paragraph()
    doc.add_paragraph()

    add_styled_table(doc,
        ["Role", "Name", "Signature", "Date"],
        [
            ["POD Lead", "Amit", "______________________", "___________"],
            ["Implementation Manager", "Shivani", "______________________", "___________"],
            ["Client Sponsor", "Prasanna", "______________________", "___________"],
        ]
    )

    doc.add_page_break()

    # ===== 1. MISSION =====
    add_heading(doc, "1. Mission", level=1)

    add_body(doc, (
        "Deliver a working AI copilot that enables support agents to resolve tickets faster and more "
        "consistently by automating classification, knowledge retrieval, action recommendation, and "
        "response drafting \u2014 targeting 85% accuracy and autonomous handling of 30% of recurring "
        "tickets, while maintaining human approval on every response."
    ))

    add_body(doc, (
        "The pilot will validate feasibility, establish baseline metrics, and produce a complete "
        "knowledge transfer package for production deployment on the client's premises."
    ))

    # ===== 2. SUCCESS CRITERIA =====
    add_heading(doc, "2. Success Criteria", level=1)

    add_styled_table(doc,
        ["#", "Metric", "Type", "Target", "Measurement"],
        [
            ["1", "Classification accuracy", "AI Quality", "\u2265 85%", "Exact match on category + priority"],
            ["2", "Retrieval accuracy", "AI Quality", "\u2265 85%", "Expected KB article in top-K set"],
            ["3", "Action accuracy", "AI Quality", "\u2265 85%", "Exact match on recommended action"],
            ["4", "Response acceptance", "Business", "\u2265 70%", "% drafts accepted without major edits"],
            ["5", "Auto-answer coverage", "Business", "30%", "% recurring tickets handled autonomously"],
        ]
    )

    doc.add_paragraph()
    add_callout(doc,
        "Go-live validation: 1,000 synthetic questions evaluated by client's support team lead "
        "before production readiness sign-off."
    )

    # ===== 3. SCOPE =====
    add_heading(doc, "3. Scope and Out-of-Scope", level=1)

    add_heading(doc, "In Scope (Phase 1 Pilot)", level=2)
    in_scope = [
        "Ticket classification (category, priority, sentiment)",
        "KB article retrieval via hybrid search (vector + BM25)",
        "Action recommendation (Reply / Ask for more info / Escalate) with reasoning",
        "Response drafting grounded in KB articles with citations",
        "Confidence scoring per output",
        "Feedback loop (agent rates/edits responses, system stores corrections)",
        "Guardrails (profanity filter, misuse prevention, graceful failure)",
        "Standalone web application (three-panel dashboard)",
        "Evaluation harness with automated scoring and 1,000-question synthetic run",
        "Complete documentation and knowledge transfer package",
    ]
    for item in in_scope:
        add_bullet(doc, item)

    add_heading(doc, "Out of Scope (Phase 1)", level=2)
    out_scope = [
        "Freshdesk API integration (architecture designed for it, not built)",
        "Multi-language support (English only)",
        "Customer-facing AI (agent-facing only)",
        "Auto-send / autonomous resolution without human approval",
        "Live KB refresh (static dataset for pilot)",
        "Production deployment, scaling, load testing",
        "Phase 2 features (not yet defined)",
    ]
    for item in out_scope:
        add_bullet(doc, item)

    # ===== 4. OPERATING CADENCE =====
    add_heading(doc, "4. Operating Cadence", level=1)

    add_styled_table(doc,
        ["Ceremony", "Frequency", "Duration", "Participants", "Owner"],
        [
            ["Sprint cycle", "2 sprints total", "Sprint 1: May 1\u201310\nSprint 2: May 11\u201316", "Full POD", "Shivani"],
            ["Weekly status email", "Weekly", "Async (written)", "Prasanna, POD", "Shivani"],
            ["Weekly sync call", "Weekly", "30 min (Google Meet)", "Prasanna, Amit, Shivani", "Shivani"],
            ["Sprint demo", "Per sprint", "30\u201345 min", "Prasanna, full POD", "Amit"],
            ["POD standup", "Daily", "15 min (internal)", "Full POD", "Amit"],
            ["Sprint retro", "Per sprint", "30 min (internal)", "Full POD", "Amit"],
        ]
    )

    doc.add_paragraph()
    add_body(doc, (
        "Channels: Email for written updates, Google Meet for calls. "
        "Client decision turnaround: 2 hours to 1 business day."
    ))

    # ===== 5. DECISION RIGHTS =====
    add_heading(doc, "5. Decision Rights", level=1)

    add_styled_table(doc,
        ["Tier", "Authority", "Examples"],
        [
            ["POD-Internal", "POD Lead + owning role",
             "Library choice, prompt structure, code style, sprint task ordering"],
            ["POD Lead", "Amit, informed by POD",
             "Architecture patterns, model selection, evaluation thresholds, release readiness"],
            ["Client Approval", "Prasanna + Shivani",
             "Scope changes, milestone shifts, data access changes, success criteria changes"],
            ["Gyde Leadership", "Engineering Director",
             "Deviation from framework non-negotiables, commercial changes"],
        ]
    )

    # ===== 6. ESCALATION PATHS =====
    add_heading(doc, "6. Escalation Paths", level=1)

    add_styled_table(doc,
        ["Type", "Gyde Contact", "Client Contact"],
        [
            ["Technical", "Amit (POD Lead)", "Prasanna"],
            ["Delivery / Commercial", "Shivani (PM)", "Prasanna"],
            ["Governance / Security", "Shubham (Governance Eng)", "Prasanna"],
            ["2nd Level (Gyde)", "Shubham (Escalation SPOC)", "\u2014"],
        ]
    )

    doc.add_paragraph()
    add_callout(doc,
        "Escalation protocol: Same-day transparency. If any risk materializes, the client hears "
        "about it within the same business day via email, not deferred to the weekly update.",
        bg="FFF8E8", border="E86C00"
    )

    # ===== 7. DEFINITION OF DONE =====
    doc.add_page_break()
    add_heading(doc, "7. Definition of Done", level=1)

    add_body(doc, "An increment is releasable to the client environment when ALL of the following are met:")

    add_styled_table(doc,
        ["#", "Gate", "Verified By"],
        [
            ["1", "All acceptance criteria for committed stories are met", "Nishka (QA)"],
            ["2", "Evaluation metrics at or above target thresholds", "Nishka + Amit"],
            ["3", "No critical or high-severity bugs open", "Nishka"],
            ["4", "Code reviewed and merged to main branch", "Amit"],
            ["5", "All prompts and data versioned in source control", "Atharva + Nancy"],
            ["6", "Security review passed (no blocking findings)", "Shubham"],
            ["7", "Documentation updated (architecture, ADRs, runbooks)", "Amit"],
            ["8", "Demo-ready in staging environment", "Amit"],
        ]
    )

    # ===== 8. RISKS AND ASSUMPTIONS =====
    add_heading(doc, "8. Risks and Assumptions", level=1)

    add_heading(doc, "Top 5 Risks", level=2)

    add_styled_table(doc,
        ["#", "Risk", "L", "I", "Mitigation", "Owner"],
        [
            ["1", "Low dataset diversity (11 unique scenarios) limits generalization",
             "H", "H", "Generate diverse synthetic data early; flag limitation", "Nishka + Atharva"],
            ["2", "Tight timeline (16 days) with hard deadlines, no buffer",
             "H", "H", "Ruthless prioritization; cut polish, not core; daily standups", "Shivani + Amit"],
            ["3", "Gemini accuracy may not reach 85% on all metrics on first pass",
             "M", "H", "LLM-agnostic architecture; fallback to GPT-4o; iterate prompts", "Atharva + Amit"],
            ["4", "Reporting category has zero training data but is in eval set",
             "M", "M", "Add 2\u20133 synthetic Reporting tickets; accept cold-start", "Nancy"],
            ["5", "On-prem handover requirements may surface late constraints",
             "L", "M", "Document all dependencies early; self-hostable components only", "Amit"],
        ]
    )

    doc.add_paragraph()

    add_heading(doc, "Assumptions", level=2)

    assumptions = [
        "Excel dataset is representative of production ticket patterns",
        "Prasanna is available for decisions within 1 business day",
        "GCP Vertex AI APIs are stable and available throughout the pilot",
        "85% accuracy is achievable with the provided KB content",
        "Team members are dedicated to this engagement (no competing priorities)",
    ]
    for i, item in enumerate(assumptions, 1):
        add_bullet(doc, item, bold_prefix=f"A{i}")

    # ===== NON-NEGOTIABLES =====
    add_heading(doc, "Framework Non-Negotiable Adherence", level=1)

    add_body(doc, "Per Doc 01, Section 5.1, this engagement adheres to all five framework non-negotiables:")

    add_styled_table(doc,
        ["#", "Non-Negotiable", "How We Fulfill It"],
        [
            ["1", "Threat modeling & secrets management",
             "Threat model delivered by Shubham; API keys via GCP Secret Manager, never in code"],
            ["2", "Evaluation before production",
             "Eval harness gates every release; 1,000-question run before go-live"],
            ["3", "Versioned data & prompts",
             "All prompts, datasets, configs in Git; every change is a tracked commit"],
            ["4", "Audit trail for AI decisions",
             "Every copilot decision logged with input, output, confidence, reasoning, sources"],
            ["5", "Incident response readiness",
             "Runbooks for top failure modes included in knowledge transfer package"],
        ]
    )

    # ===== POD COMPOSITION =====
    doc.add_paragraph()
    add_heading(doc, "POD Composition", level=1)

    add_styled_table(doc,
        ["Role", "Name", "Key Responsibilities"],
        [
            ["POD Lead", "Amit", "Architecture, UI, code review, tech decisions, demos"],
            ["AI Engineer", "Atharva", "LLM prompts, retrieval pipeline, confidence scoring, feedback loop"],
            ["Data Engineer", "Nancy", "Data ingestion, KB indexing, embeddings, vector store, data quality"],
            ["QA", "Nishka", "Eval harness, golden dataset, synthetic data, adversarial testing"],
            ["Governance Engineer", "Shubham", "Threat model, guardrails, security review, compliance"],
            ["Implementation Manager", "Shivani", "Charter, sprint planning, status reports, risk register, client comms"],
        ]
    )

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "\u2014  End of Charter  \u2014", size=10, color=MEDIUM_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, (
        "This charter is effective upon signature by all three signatories and remains in force "
        "for the duration of the engagement. Any material changes require written agreement from all parties."
    ), size=9, color=MEDIUM_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "Gyde AI POD  |  Confidential  |  2026-04-30", size=9, color=MEDIUM_GRAY)

    return doc


if __name__ == "__main__":
    doc = build_doc()

    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "client-docs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "POD_Charter_AI_Support_Copilot.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")

    sim_dir = "/Users/amit/Work/Gyde/Gyde Pivot/Gyde AI POD Framework/Simulation Docs"
    sim_path = os.path.join(sim_dir, "POD_Charter_AI_Support_Copilot.docx")
    doc.save(sim_path)
    print(f"Saved: {sim_path}")
