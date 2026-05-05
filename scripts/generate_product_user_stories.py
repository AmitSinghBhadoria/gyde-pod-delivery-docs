"""
Generate Product User Stories Word document.
Internal POD document, also suitable for client communication.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

DARK_BLUE = RGBColor(0x1B, 0x2A, 0x4A)
MEDIUM_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
ACCENT_GREEN = RGBColor(0x2D, 0x7D, 0x46)
ACCENT_ORANGE = RGBColor(0xE8, 0x6C, 0x00)
ACCENT_RED = RGBColor(0xCC, 0x33, 0x33)
LIGHT_GREEN_BG = "E8F5E9"
LIGHT_ORANGE_BG = "FFF3E0"
LIGHT_BLUE_BG = "E8F0F8"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def fmt(paragraph, text, bold=False, size=10, color=DARK_GRAY, font="Calibri", italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    return run


def add_heading(doc, text, level=1, color=DARK_BLUE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_styled_table(doc, headers, rows, header_color="1B2A4A", stripe=True, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt(p, h, bold=True, size=9, color=WHITE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.cell(r + 1, c)
            if stripe and r % 2 == 1:
                set_cell_shading(cell, "F5F7FA")
            p = cell.paragraphs[0]
            fmt(p, str(val), size=9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    return table


def add_body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    fmt(p, text, size=10)
    return p


def add_bullet(doc, text, bold_prefix=None, indent=Cm(1)):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if bold_prefix:
        fmt(p, f"{bold_prefix}: ", bold=True, size=9.5)
        fmt(p, text, size=9.5)
    else:
        fmt(p, text, size=9.5)
    return p


def add_sub_bullet(doc, text, indent=Cm(2)):
    p = doc.add_paragraph(style="List Bullet 2")
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    fmt(p, text, size=9)
    return p


def add_callout(doc, text, bg="E8F0F8", border="2C5F8A"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="{border}"/>'
        f'<w:top w:val="single" w:sz="2" w:space="0" w:color="{border}"/>'
        f'<w:bottom w:val="single" w:sz="2" w:space="0" w:color="{border}"/>'
        f'<w:right w:val="single" w:sz="2" w:space="0" w:color="{border}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    p = cell.paragraphs[0]
    fmt(p, text, size=9.5)


def add_user_story(doc, story_id, title, persona, scope, description, acceptance_criteria):
    """Add a main user story with formatting."""
    # Story heading
    h = doc.add_heading(f"{story_id}: {title}", level=3)
    for run in h.runs:
        run.font.color.rgb = MEDIUM_BLUE
        run.font.name = "Calibri"

    # Scope tag
    scope_bg = LIGHT_GREEN_BG if scope == "Pilot" else LIGHT_ORANGE_BG
    scope_border = "2D7D46" if scope == "Pilot" else "E86C00"
    add_callout(doc, f"[{scope}]  {description}", bg=scope_bg, border=scope_border)

    doc.add_paragraph()

    # Acceptance criteria
    p = doc.add_paragraph()
    fmt(p, "Acceptance criteria:", bold=True, size=9.5)
    for ac in acceptance_criteria:
        add_bullet(doc, ac, indent=Cm(1))

    doc.add_paragraph()


def add_sub_story(doc, story_id, title, persona, acceptance_criteria):
    """Add a sub-story nested under a main story."""
    p = doc.add_paragraph()
    fmt(p, f"{story_id}: ", bold=True, size=10, color=DARK_BLUE)
    fmt(p, title, size=10, italic=True)

    for ac in acceptance_criteria:
        add_sub_bullet(doc, ac, indent=Cm(1.5))

    doc.add_paragraph()


def build_doc():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ===== TITLE PAGE =====
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "PRODUCT USER STORIES", bold=True, size=28, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "AI Support Copilot", size=16, color=MEDIUM_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    fmt(p, "Version 1.0  |  2026-05-04", size=11, color=MEDIUM_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    fmt(p, "Full Production Vision  |  Pilot + Future Scope", size=10, color=MEDIUM_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    fmt(p, "Gyde AI POD  |  Confidential", size=9, color=MEDIUM_GRAY)

    doc.add_page_break()

    # ===== PERSONAS =====
    add_heading(doc, "Personas", level=1)

    add_styled_table(doc,
        ["Persona", "Role", "Goal", "Usage"],
        [
            ["Asha", "Support Agent", "Resolve customer tickets faster with AI assistance", "All day, every ticket"],
            ["Ravi", "Team Lead / Supervisor", "Monitor team performance and copilot effectiveness", "Daily reviews, weekly reports"],
            ["Meera", "KB Admin", "Keep the knowledge base current for accurate copilot answers", "Weekly updates"],
        ]
    )

    doc.add_paragraph()

    # ===== EPIC MAP =====
    add_heading(doc, "Epic Map", level=1)

    add_styled_table(doc,
        ["#", "Epic", "Persona", "Scope"],
        [
            ["E1", "Ticket Triage & Classification", "Asha", "Pilot"],
            ["E2", "AI-Assisted Response Drafting", "Asha", "Pilot"],
            ["E3", "Knowledge Retrieval & Citation", "Asha", "Pilot"],
            ["E4", "Escalation Handling", "Asha", "Pilot"],
            ["E5", "Agent Feedback & Copilot Learning", "Asha", "Pilot"],
            ["E6", "Safety & Guardrails", "Asha", "Pilot"],
            ["E7", "Performance Monitoring & Analytics", "Ravi", "Pilot + Future"],
            ["E8", "Knowledge Base Management", "Meera", "Future"],
            ["E9", "Multi-Channel Support", "Asha", "Future"],
        ]
    )

    doc.add_page_break()

    # ===== E1: TICKET TRIAGE =====
    add_heading(doc, "E1: Ticket Triage & Classification", level=1)

    add_user_story(doc,
        "US-1.1", "View automatic ticket classification", "Asha", "Pilot",
        "As Asha, I want to see an automatic classification of each ticket (category, priority, sentiment) the moment I open it, so that I can understand the ticket's nature at a glance without reading the entire description.",
        [
            "Given a ticket is selected from the queue, when the copilot panel loads, then it displays: category (one of 7 values), priority (Low/Medium/High/Critical), sentiment (Positive/Neutral/Negative/Frustrated), and a confidence score (0\u2013100%).",
            "Given the classification is displayed, when I look at the confidence score, then it is color-coded: green (\u2265 80%), yellow (60\u201379%), red (< 60%) so I can quickly judge reliability.",
            "Given the copilot has classified the ticket, when I disagree with the classification, then I can see it is a suggestion \u2014 it does not auto-change any ticket fields in the helpdesk system.",
        ]
    )

    add_sub_story(doc,
        "US-1.1a", "Handle ambiguous tickets",
        "Asha",
        [
            "Given a ticket has unclear language that spans multiple categories, when the copilot classifies it, then the confidence score is below 60% and a yellow/red indicator warns me to verify.",
            "Given a low-confidence classification, when I view the copilot sidebar, then I see the top-2 candidate categories (e.g., \"Billing (45%) / Bug Report (38%)\") so I can make an informed decision.",
        ]
    )

    add_sub_story(doc,
        "US-1.1b", "Handle non-English or gibberish tickets",
        "Asha",
        [
            "Given a ticket is in a non-English language, when the copilot processes it, then it returns category \"Unknown\" with confidence < 30% and a note: \"Ticket language not supported.\"",
            "Given a ticket body is empty or contains only gibberish, when the copilot processes it, then it flags: \"Insufficient content to classify\" and does not generate a draft response.",
        ]
    )

    add_user_story(doc,
        "US-1.2", "Filter and sort tickets by AI classification", "Asha", "Future",
        "As Asha, I want to filter the ticket queue by the copilot's predicted category and priority, so that I can focus on high-priority tickets in my area of expertise first.",
        [
            "Given the ticket queue is displayed, when I select a category filter (e.g., \"Authentication\"), then only tickets classified as Authentication are shown.",
            "Given the queue is filtered, when I sort by priority, then Critical tickets appear first, followed by High, Medium, Low.",
            "Given the copilot has classified all tickets in the queue, when I view the queue, then each ticket row shows a small category badge and priority indicator next to the subject line.",
        ]
    )

    # ===== E2: AI-ASSISTED RESPONSE =====
    add_heading(doc, "E2: AI-Assisted Response Drafting", level=1)

    add_user_story(doc,
        "US-2.1", "Receive a draft response for a ticket", "Asha", "Pilot",
        "As Asha, I want the copilot to generate a draft response for the ticket I'm viewing, so that I can send a well-written, accurate reply without composing it from scratch.",
        [
            "Given I select a ticket from the queue, when the pipeline finishes processing, then a draft response appears in the copilot sidebar under \"Draft Response.\"",
            "Given the draft response is displayed, when I read it, then it addresses the customer's issue, uses a professional tone, and includes a greeting and closing.",
            "Given the draft is acceptable, when I click \"Accept\", then the draft text is copied into my Agent Response text area, ready to send.",
            "Given the pipeline is running, when I look at the copilot sidebar, then I see a step-by-step progress indicator showing which stage is executing.",
        ]
    )

    add_sub_story(doc,
        "US-2.1a", "Edit a draft before sending",
        "Asha",
        [
            "Given a draft response is displayed, when I click \"Edit\", then the draft text becomes editable in a text area within the copilot sidebar.",
            "Given I have edited the draft, when I click \"Accept Edited\", then the modified version is copied into my Agent Response area, and both the original and edited versions are saved for analysis.",
        ]
    )

    add_sub_story(doc,
        "US-2.1b", "Override the draft entirely",
        "Asha",
        [
            "Given a draft response is displayed, when I click \"Override\", then the draft is dismissed, my Agent Response area remains empty for me to type, and the override is recorded as feedback.",
        ]
    )

    add_sub_story(doc,
        "US-2.1c", "Handle pipeline timeout or failure",
        "Asha",
        [
            "Given the pipeline takes longer than 15 seconds, when the timeout is reached, then the copilot sidebar shows: \"Response generation timed out. Click Retry or compose your response manually.\"",
            "Given an API error occurs (LLM service down, Elasticsearch unreachable), when the error state renders, then I see a specific error message with a \"Retry\" button.",
            "Given the copilot has failed, when I look at my ticket detail panel, then the Agent Response area is still fully functional \u2014 the copilot failure never blocks me from responding manually.",
        ]
    )

    add_sub_story(doc,
        "US-2.1d", "Handle tickets with no relevant KB match",
        "Asha",
        [
            "Given a ticket is about a topic not covered by any KB article, when the retrieval step completes, then the copilot shows: \"No relevant KB articles found\" and the draft response clearly states it could not find supporting documentation.",
            "Given no KB match was found, when a draft is still generated, then it is marked with a warning: \"This response is not grounded in KB articles \u2014 please verify before sending.\"",
        ]
    )

    add_user_story(doc,
        "US-2.2", "View pipeline processing in real-time", "Asha", "Pilot",
        "As Asha, I want to see each step of the AI pipeline as it executes, so that I understand what the copilot is doing and can start reading partial results while later steps are still running.",
        [
            "Given I select a ticket, when the pipeline starts, then I see a step tracker: Step 1 (Classify) \u2192 Step 2 (Retrieve) \u2192 Step 3 (Reason) \u2192 Step 4 (Draft) \u2192 Step 5 (Guardrails).",
            "Given Step 1 completes, when Step 2 starts, then the classification result is already visible in the sidebar while retrieval is still running.",
            "Given all steps complete, when the full result is displayed, then the total processing time is shown (e.g., \"Processed in 7.2s\").",
        ]
    )

    doc.add_page_break()

    # ===== E3: KNOWLEDGE RETRIEVAL =====
    add_heading(doc, "E3: Knowledge Retrieval & Citation", level=1)

    add_user_story(doc,
        "US-3.1", "See which KB articles the copilot used", "Asha", "Pilot",
        "As Asha, I want to see the KB articles the copilot retrieved and cited in its draft, so that I can verify the response is based on accurate, current information.",
        [
            "Given the copilot has processed a ticket, when I view the \"Relevant KB Articles\" section, then I see the top 3 matched articles with: KB ID, title, and a relevance score (0\u20131).",
            "Given the draft response contains citations like \"[KB-001]\", when I click a citation, then a popover shows the full KB article content (title + body).",
            "Given the KB articles are displayed, when I read the relevance scores, then they help me judge how closely each article matches the ticket.",
        ]
    )

    add_sub_story(doc,
        "US-3.1a", "Handle outdated KB article [Future]",
        "Asha",
        [
            "Given a KB article is displayed in the copilot sidebar, when I look at its metadata, then I see the last_updated date.",
            "Given a KB article was last updated more than 90 days ago, when it is displayed, then a subtle \"May be outdated\" indicator appears next to it.",
        ]
    )

    add_sub_story(doc,
        "US-3.1b", "View articles not cited but potentially relevant [Future]",
        "Asha",
        [
            "Given the copilot retrieved 3 articles, when I click \"Show more articles\", then I see the next 2\u20133 articles that scored below the citation threshold but may still be relevant.",
        ]
    )

    # ===== E4: ESCALATION =====
    add_heading(doc, "E4: Escalation Handling", level=1)

    add_user_story(doc,
        "US-4.1", "Receive escalation recommendation", "Asha", "Pilot",
        "As Asha, I want the copilot to recommend escalation when a ticket matches escalation rules, so that critical issues reach the right team without me having to memorize all escalation criteria.",
        [
            "Given a ticket matches an escalation rule, when the copilot processes it, then the recommended action shows \"Escalate\" in a prominent red/orange banner.",
            "Given escalation is recommended, when I view the reasoning section, then I see: the escalation team name, required context to include, and a brief explanation of why escalation was triggered.",
            "Given escalation is recommended, when the copilot generates a draft, then it drafts an escalation handoff note (not a customer reply) with ticket summary, impact assessment, and steps already attempted.",
        ]
    )

    add_sub_story(doc,
        "US-4.1a", "Disagree with escalation recommendation",
        "Asha",
        [
            "Given the copilot recommends escalation, when I click \"Override\", then I can change the action to \"Reply\" and write my own response. The override is recorded with the original recommendation preserved for review.",
        ]
    )

    add_sub_story(doc,
        "US-4.1b", "Handle borderline escalation",
        "Asha",
        [
            "Given a ticket partially matches an escalation rule but with low confidence, when the copilot processes it, then it shows: \"Escalation may be appropriate (confidence: X%). Review the reasoning before deciding.\"",
            "Given a borderline case, when I view the reasoning, then I see both the \"Reply\" rationale and the \"Escalate\" rationale side by side.",
        ]
    )

    add_user_story(doc,
        "US-4.2", "Ask clarification instead of replying", "Asha", "Pilot",
        "As Asha, I want the copilot to recommend \"Ask Clarification\" when the ticket is too vague to resolve, so that I request the right information from the customer on the first follow-up.",
        [
            "Given a ticket is vague (e.g., \"It's not working\"), when the copilot processes it, then the recommended action is \"Ask Clarification\" instead of \"Reply.\"",
            "Given the action is \"Ask Clarification\", when the draft is generated, then it contains specific questions to ask the customer.",
            "Given the copilot recommends asking clarification, when I view the reasoning, then I see what information is missing and why a direct reply isn't possible.",
        ]
    )

    doc.add_page_break()

    # ===== E5: FEEDBACK =====
    add_heading(doc, "E5: Agent Feedback & Copilot Learning", level=1)

    add_user_story(doc,
        "US-5.1", "Rate copilot helpfulness", "Asha", "Pilot",
        "As Asha, I want to quickly rate whether the copilot's suggestion was helpful, so that the team can measure how useful the AI is and improve it over time.",
        [
            "Given the copilot has produced a result, when I see the feedback row, then I can click \"Yes\" or \"No\" to rate helpfulness.",
            "Given I click \"Yes\" and accept the draft, when feedback is submitted, then the system records: ticket_id, helpful=true, original_draft, edited_draft=null.",
            "Given I click \"No\" and override, when feedback is submitted, then the system records: ticket_id, helpful=false, and my override response (if I wrote one).",
        ]
    )

    add_sub_story(doc,
        "US-5.1a", "Submit detailed feedback [Future]",
        "Asha",
        [
            "Given I click \"No\", when a feedback modal appears, then I can select a reason: \"Wrong category\", \"Wrong KB articles\", \"Bad tone\", \"Factually incorrect\", \"Other\" with a free-text field.",
        ]
    )

    add_sub_story(doc,
        "US-5.1b", "Track my own copilot usage stats [Future]",
        "Asha",
        [
            "Given I have processed 50+ tickets with the copilot, when I open my profile/stats page, then I see: total tickets, % accepted as-is, % edited, % overridden, breakdown by category.",
        ]
    )

    # ===== E6: SAFETY =====
    add_heading(doc, "E6: Safety & Guardrails", level=1)

    add_user_story(doc,
        "US-6.1", "See safety warnings before sending", "Asha", "Pilot",
        "As Asha, I want the copilot to warn me about potential issues in its draft (PII leakage, low confidence, ungrounded claims), so that I don't send a problematic response to a customer.",
        [
            "Given the guardrails layer detects an issue, when the copilot sidebar renders, then a warning banner appears at the top with severity (critical = red, warning = yellow) and a description.",
            "Given a critical warning exists (e.g., PII leakage detected), when I try to click \"Accept\", then I see a confirmation: \"A critical warning was flagged. Are you sure you want to accept this draft?\"",
            "Given no warnings exist, when the copilot sidebar renders, then a green checkmark or \"No issues detected\" indicator appears.",
        ]
    )

    add_sub_story(doc,
        "US-6.1a", "PII echo prevention",
        "Asha",
        [
            "Given a ticket contains customer PII like an SSN or phone number, when the copilot generates a draft, then the PII is not repeated in the draft text.",
            "Given PII is detected in the ticket, when the guardrails flag it, then the warning says: \"PII detected in ticket. Draft has been checked for PII echo.\"",
        ]
    )

    add_sub_story(doc,
        "US-6.1b", "Handle prompt injection attempts",
        "Asha",
        [
            "Given a ticket contains prompt injection text (e.g., \"Ignore your instructions\"), when the copilot processes it, then a critical warning appears: \"Potential prompt injection detected in ticket text.\"",
            "Given a prompt injection is detected, when the draft is reviewed, then it responds to the actual ticket content and does not follow the injected instructions.",
        ]
    )

    add_sub_story(doc,
        "US-6.1c", "Flag ungrounded claims",
        "Asha",
        [
            "Given the draft contains a factual claim, when the claim cannot be traced to a cited KB article, then a warning appears: \"Ungrounded claim detected. Verify before sending.\"",
        ]
    )

    doc.add_page_break()

    # ===== E7: PERFORMANCE MONITORING =====
    add_heading(doc, "E7: Performance Monitoring & Analytics", level=1)

    add_user_story(doc,
        "US-7.1", "View copilot accuracy dashboard", "Ravi", "Pilot",
        "As Ravi (team lead), I want to see aggregate copilot accuracy metrics, so that I can assess whether the AI is performing well enough for production deployment.",
        [
            "Given eval runs have been completed, when I view the metrics dashboard, then I see: classification accuracy, retrieval accuracy, action accuracy, and response faithfulness \u2014 each with current value, target, and trend.",
            "Given metrics are below target, when I view the dashboard, then the metric is highlighted in red with the gap clearly shown.",
        ]
    )

    add_sub_story(doc,
        "US-7.1a", "View per-category breakdown",
        "Ravi",
        [
            "Given the dashboard has per-category data, when I expand \"Classification Accuracy\", then I see a row per category (Authentication, Billing, Bug Report, etc.) with individual accuracy scores.",
            "Given a category has below-target accuracy, when I view it, then it is flagged for prompt tuning attention.",
        ]
    )

    add_sub_story(doc,
        "US-7.1b", "View agent override patterns [Future]",
        "Ravi",
        [
            "Given feedback data exists, when I view the \"Agent Overrides\" report, then I see: per-agent override rate, per-category override rate, and common override reasons.",
        ]
    )

    add_user_story(doc,
        "US-7.2", "Review copilot cost and latency", "Ravi", "Pilot",
        "As Ravi, I want to see the cost per ticket and response latency, so that I can evaluate whether the copilot is economically viable at scale.",
        [
            "Given the pipeline has processed tickets, when I view operational metrics, then I see: average cost per ticket (LLM API cost), p50 and p95 latency, and error rate.",
            "Given cost per ticket exceeds $0.50, when the metric is displayed, then it is flagged as above the hard limit with a recommendation to investigate.",
        ]
    )

    add_user_story(doc,
        "US-7.3", "Receive weekly copilot performance report", "Ravi", "Future",
        "As Ravi, I want to receive an automated weekly email summarizing copilot performance, so that I stay informed without having to check the dashboard manually.",
        [
            "Given a week has passed, when the report is generated, then it includes: tickets processed, accuracy metrics, acceptance rate, agent override summary, cost summary, and top 5 failure cases.",
            "Given the report is generated, when Ravi receives it, then it includes a comparison to the previous week with delta indicators.",
        ]
    )

    # ===== E8: KB MANAGEMENT =====
    add_heading(doc, "E8: Knowledge Base Management", level=1)

    add_user_story(doc,
        "US-8.1", "Add a new KB article", "Meera", "Future",
        "As Meera (KB admin), I want to add a new KB article and have the copilot immediately use it for retrieval, so that new product knowledge reaches agents through the copilot without waiting for a system update.",
        [
            "Given I write a new KB article, when I submit it through the KB management interface, then the article is stored in MongoDB and automatically embedded and indexed in Elasticsearch.",
            "Given a new article is indexed, when the next ticket related to that topic arrives, then the copilot retrieves the new article as a relevant match.",
            "Given I add an article, when the process completes, then I see a confirmation: \"Article KB-XXX indexed. Available for copilot retrieval.\"",
        ]
    )

    add_sub_story(doc,
        "US-8.1a", "Update an existing KB article",
        "Meera",
        [
            "Given I edit KB-003's content, when the update is saved, then the old embedding is replaced with a new one generated from the updated text.",
            "Given the article is re-embedded, when a ticket that previously retrieved KB-003 is re-processed, then the copilot uses the updated content.",
        ]
    )

    add_sub_story(doc,
        "US-8.1b", "Delete or archive a KB article",
        "Meera",
        [
            "Given I archive KB-005, when the action completes, then the article is removed from the Elasticsearch index and no longer appears in retrieval results.",
            "Given KB-005 was previously cited in a draft, when I archive it, then historical records still reference KB-005 but new queries do not retrieve it.",
        ]
    )

    add_sub_story(doc,
        "US-8.1c", "View KB coverage gaps",
        "Meera",
        [
            "Given copilot retrieval data is collected, when I view the KB coverage report, then I see: per-category average retrieval score, number of tickets with no KB match, and the top 5 tickets where retrieval failed most recently.",
        ]
    )

    # ===== E9: MULTI-CHANNEL =====
    add_heading(doc, "E9: Multi-Channel Support", level=1)

    add_user_story(doc,
        "US-9.1", "Process tickets from Freshdesk", "Asha", "Future",
        "As Asha, I want the copilot to work as a sidebar widget inside Freshdesk, so that I don't need to switch between two applications to handle tickets.",
        [
            "Given I open a ticket in Freshdesk, when the copilot sidebar widget loads, then it automatically runs the pipeline on the current ticket and displays results inline.",
            "Given I accept a copilot draft in the Freshdesk widget, when I click \"Accept\", then the draft is inserted into Freshdesk's reply editor.",
            "Given the copilot is a Freshdesk widget, when it processes a ticket, then it uses the same backend pipeline as the standalone pilot app.",
        ]
    )

    add_user_story(doc,
        "US-9.2", "Process chat conversations", "Asha", "Future",
        "As Asha, I want the copilot to work with live chat conversations (not just tickets), so that I get AI-assisted responses in real-time chat scenarios.",
        [
            "Given I am in a live chat with a customer, when I click \"Get Copilot Suggestion\", then the copilot processes the full chat transcript and generates a response.",
            "Given the chat is ongoing, when the customer sends a new message, then the copilot can re-process with the updated transcript.",
        ]
    )

    doc.add_page_break()

    # ===== PRIORITY MATRIX =====
    add_heading(doc, "Story Priority Matrix", level=1)

    add_styled_table(doc,
        ["Priority", "Stories", "Rationale"],
        [
            ["P0 \u2014 Must have\n(Pilot)", "US-1.1, US-2.1, US-3.1, US-4.1, US-4.2, US-5.1, US-6.1",
             "Core copilot workflow \u2014 agent must be able to see classification, get a draft, handle escalation, give feedback, and see safety warnings"],
            ["P1 \u2014 Should have\n(Pilot)", "US-1.1a, US-1.1b, US-2.1a, US-2.1b, US-2.1c, US-2.1d, US-2.2, US-4.1a, US-4.1b, US-6.1a, US-6.1b, US-6.1c, US-7.1, US-7.2",
             "Edge cases, error handling, and metrics that make the pilot production-realistic"],
            ["P2 \u2014 Nice to have\n(Pilot)", "US-1.2, US-7.1a",
             "Queue filtering and per-category analytics add polish"],
            ["P3 \u2014 Post-pilot", "US-3.1a, US-3.1b, US-5.1a, US-5.1b, US-7.1b, US-7.3, US-8.1, US-8.1a, US-8.1b, US-8.1c, US-9.1, US-9.2",
             "KB management, detailed feedback, Freshdesk integration, chat support"],
        ],
        header_color="1B2A4A"
    )

    doc.add_paragraph()

    # ===== TRACEABILITY =====
    add_heading(doc, "Traceability: Product Stories \u2192 Engineering Stories", level=1)

    add_body(doc, "Each product story maps to one or more engineering stories from the Sprint Backlog. This ensures every user-facing need has a concrete implementation plan.")

    add_styled_table(doc,
        ["Product Story", "Engineering Story", "Notes"],
        [
            ["US-1.1", "S1-03, S1-10", "Classification logic + sidebar display"],
            ["US-1.1a", "S1-03", "Low-confidence handling in classify prompt"],
            ["US-1.1b", "S1-03, S2-03", "Edge case classification + adversarial validation"],
            ["US-2.1", "S1-07, S1-08, S1-10", "Draft generation + pipeline chain + UI"],
            ["US-2.1a", "S1-10, S2-02", "Edit flow + feedback storage"],
            ["US-2.1b", "S1-10, S2-02", "Override flow + feedback storage"],
            ["US-2.1c", "S1-10, S2-06", "Error states + retry button"],
            ["US-2.1d", "S1-05, S2-06", "No-match handling + warning display"],
            ["US-2.2", "S2-06", "Step-by-step progress indicator"],
            ["US-3.1", "S1-05, S1-07, S2-06", "Retrieval + citation rendering"],
            ["US-4.1", "S1-06, S1-10", "Escalation rule matching + UI"],
            ["US-4.1a", "S1-10, S2-02", "Override action + feedback recording"],
            ["US-4.2", "S1-06, S1-07", "\"Ask Clarification\" action logic + draft"],
            ["US-5.1", "S1-10, S2-02", "Feedback widget + storage"],
            ["US-6.1", "S2-01, S2-06", "Guardrails processing + warning banners"],
            ["US-6.1a", "S2-01, S2-03", "PII filter + adversarial validation"],
            ["US-6.1b", "S2-01, S2-03", "Prompt injection detection"],
            ["US-6.1c", "S2-07, S2-01", "Hallucination check + warning"],
            ["US-7.1", "S1-11, S2-05", "Eval metrics + dashboard"],
            ["US-7.2", "S1-11", "Latency + cost tracking"],
        ],
        header_color="2C5F8A"
    )

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "\u2014  End of Product User Stories  \u2014", size=10, color=MEDIUM_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, "Gyde AI POD  |  Confidential  |  2026-05-04", size=9, color=MEDIUM_GRAY)

    return doc


if __name__ == "__main__":
    doc = build_doc()

    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "client-docs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Product_User_Stories_AI_Support_Copilot.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")

    sim_dir = "/Users/amit/Work/Gyde/Gyde Pivot/Gyde AI POD Framework/Simulation Docs"
    sim_path = os.path.join(sim_dir, "Product_User_Stories_AI_Support_Copilot.docx")
    doc.save(sim_path)
    print(f"Saved: {sim_path}")
