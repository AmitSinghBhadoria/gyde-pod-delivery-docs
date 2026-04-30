"""
Generate Use Case Canvas Word document in business model canvas style.
Professional layout for sharing with clients.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# --- Color palette ---
DARK_BLUE = RGBColor(0x1B, 0x2A, 0x4A)
MEDIUM_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
LIGHT_BLUE = RGBColor(0xE8, 0xF0, 0xF8)
ACCENT_ORANGE = RGBColor(0xE8, 0x6C, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
BORDER_GRAY = RGBColor(0xCC, 0xCC, 0xCC)


def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            border = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{val["sz"]}" '
                f'w:space="0" w:color="{val["color"]}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_formatted_text(paragraph, text, bold=False, size=9, color=DARK_GRAY, font_name="Calibri"):
    """Add a run of formatted text to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font_name
    return run


def add_section_title(cell, title, color=MEDIUM_BLUE):
    """Add a section title to a cell."""
    p = cell.paragraphs[0]
    p.space_before = Pt(2)
    p.space_after = Pt(4)
    add_formatted_text(p, title, bold=True, size=10, color=color)


def add_bullet_items(cell, items, bold_prefix=True):
    """Add bullet-point items to a cell."""
    for item in items:
        p = cell.add_paragraph()
        p.space_before = Pt(1)
        p.space_after = Pt(1)
        p.paragraph_format.left_indent = Pt(8)
        if isinstance(item, tuple) and bold_prefix:
            add_formatted_text(p, f"{item[0]}: ", bold=True, size=8)
            add_formatted_text(p, item[1], size=8)
        else:
            add_formatted_text(p, f"  {item}", size=8)


def add_mini_table(cell, headers, rows):
    """Add a small formatted table inside a cell."""
    # Use paragraphs with tab stops to simulate a table
    for row in rows:
        p = cell.add_paragraph()
        p.space_before = Pt(1)
        p.space_after = Pt(1)
        if isinstance(row, tuple):
            add_formatted_text(p, f"  {row[0]}: ", bold=True, size=8)
            add_formatted_text(p, row[1], size=8)
        else:
            add_formatted_text(p, f"  {row}", size=8)


def set_cell_width(cell, width):
    """Set the width of a cell."""
    cell.width = width
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(width.emu / 635)}" w:type="dxa"/>')
    # Remove existing tcW
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcPr.append(tcW)


def merge_cells_in_row(table, row_idx, start_col, end_col):
    """Merge cells in a row."""
    cell = table.cell(row_idx, start_col)
    for col in range(start_col + 1, end_col + 1):
        cell = cell.merge(table.cell(row_idx, col))
    return cell


def build_canvas():
    doc = Document()

    # --- Page setup: Landscape A4 ---
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # --- Header bar ---
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = True

    # Left: Title
    cell_left = header_table.cell(0, 0)
    set_cell_shading(cell_left, "1B2A4A")
    p = cell_left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_formatted_text(p, "USE CASE CANVAS", bold=True, size=16, color=WHITE)
    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Middle: Product name
    cell_mid = header_table.cell(0, 1)
    set_cell_shading(cell_mid, "1B2A4A")
    p = cell_mid.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "AI Support Copilot", bold=True, size=14, color=RGBColor(0xE8, 0xF0, 0xF8))
    cell_mid.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Right: Meta
    cell_right = header_table.cell(0, 2)
    set_cell_shading(cell_right, "1B2A4A")
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_formatted_text(p, "Phase 1 Pilot  |  v1.0  |  2026-04-30", size=9, color=RGBColor(0xBB, 0xCC, 0xDD))
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Remove borders from header
    for cell in header_table.rows[0].cells:
        set_cell_border(cell,
            top={"sz": "0", "color": "1B2A4A"},
            bottom={"sz": "0", "color": "1B2A4A"},
            left={"sz": "0", "color": "1B2A4A"},
            right={"sz": "0", "color": "1B2A4A"})

    # tight spacer
    spacer0 = doc.add_paragraph()
    spacer0.space_before = Pt(2)
    spacer0.space_after = Pt(2)

    # --- Mission Statement (full width) ---
    mission_table = doc.add_table(rows=1, cols=1)
    mission_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = mission_table.cell(0, 0)
    set_cell_shading(cell, "E8F0F8")
    set_cell_border(cell,
        top={"sz": "6", "color": "2C5F8A"},
        bottom={"sz": "6", "color": "2C5F8A"},
        left={"sz": "6", "color": "2C5F8A"},
        right={"sz": "6", "color": "2C5F8A"})

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_formatted_text(p, "MISSION  ", bold=True, size=10, color=MEDIUM_BLUE)
    add_formatted_text(
        p,
        "Enable support agents to resolve tickets faster and more consistently by providing an AI copilot "
        "that classifies incoming tickets, retrieves relevant KB articles, recommends the next-best-action, "
        "and drafts grounded responses \u2014 reducing resolution time from hours/days to minutes while "
        "maintaining human approval on every response.",
        size=9.5, color=DARK_GRAY
    )

    # tight spacer
    spacer = doc.add_paragraph()
    spacer.space_before = Pt(2)
    spacer.space_after = Pt(2)

    # --- Main Canvas Grid: 4 columns x 3 rows ---
    canvas = doc.add_table(rows=3, cols=4)
    canvas.alignment = WD_TABLE_ALIGNMENT.CENTER

    border_style = {"sz": "4", "color": "CCCCCC"}

    # Style all cells
    for row in canvas.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top=border_style, bottom=border_style,
                left=border_style, right=border_style)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ============================================
    # ROW 0: Users & Volumes | Inputs | Outputs | Success Metrics
    # ============================================

    # --- Users & Volumes ---
    cell = canvas.cell(0, 0)
    set_cell_shading(cell, "F8F9FA")
    add_section_title(cell, "USERS & VOLUMES")
    add_bullet_items(cell, [
        ("Primary users", "Support agents (Freshdesk)"),
        ("Secondary", "Support team lead (evaluator)"),
        ("Volume", "~1,000 tickets/month"),
        ("Auto-answer target", "30% of recurring tickets (~300/mo)"),
        ("Peak estimate", "~50 tickets/day avg, peak 2-3x"),
    ])

    # --- Inputs ---
    cell = canvas.cell(0, 1)
    set_cell_shading(cell, "F8F9FA")
    add_section_title(cell, "INPUTS")
    add_bullet_items(cell, [
        ("Support ticket", "Subject + description + channel + customer"),
        ("KB articles", "12 articles across 7 categories"),
        ("Escalation rules", "5 rules mapping conditions to teams"),
        ("Source (pilot)", "Excel dataset"),
        ("Source (prod)", "Freshworks APIs"),
    ])

    # --- Outputs ---
    cell = canvas.cell(0, 2)
    set_cell_shading(cell, "F8F9FA")
    add_section_title(cell, "OUTPUTS")
    add_bullet_items(cell, [
        ("Classification", "Category + Priority + Sentiment"),
        ("Retrieved KB", "Top-k articles with relevance scores"),
        ("Action", "Reply / Ask for info / Escalate"),
        ("Draft response", "Grounded text with KB citations"),
        ("Confidence", "Score per output component"),
        ("Reasoning", "Why this recommendation was made"),
    ])

    # --- Success Metrics ---
    cell = canvas.cell(0, 3)
    set_cell_shading(cell, "FFF8E8")
    add_section_title(cell, "SUCCESS METRICS", color=ACCENT_ORANGE)
    add_bullet_items(cell, [
        ("Classification", ">= 85% accuracy"),
        ("Retrieval", ">= 85% accuracy"),
        ("Action", ">= 85% accuracy"),
        ("Response acceptance", ">= 70% without major edits"),
        ("Auto-answer", "30% of recurring tickets"),
        ("Validation", "1000 synthetic questions"),
    ])

    # ============================================
    # ROW 1: Hard Limits (2 cols) | Data Sources (2 cols)
    # ============================================

    # --- Hard Limits ---
    cell = merge_cells_in_row(canvas, 1, 0, 1)
    set_cell_shading(cell, "FFF0F0")
    add_section_title(cell, "HARD LIMITS & CONSTRAINTS", color=RGBColor(0xCC, 0x33, 0x33))
    add_bullet_items(cell, [
        ("Human-in-the-loop", "Always \u2014 copilot suggests, agent decides, no auto-send"),
        ("Profanity / misuse", "Zero tolerance \u2014 guardrails on all outputs"),
        ("LLM agnostic", "Architecture must support swapping providers"),
        ("Portability", "Full on-premises handover post-pilot (code, models, data)"),
        ("Hallucination", "Responses grounded in KB only \u2014 cite sources or flag 'no match'"),
        ("Timeline", "Demo: May 10 | Final delivery: May 16 (HARD deadlines)"),
    ])

    # --- Data Sources ---
    cell = merge_cells_in_row(canvas, 1, 2, 3)
    set_cell_shading(cell, "F0F8F0")
    add_section_title(cell, "DATA SOURCES", color=RGBColor(0x2D, 0x7D, 0x46))
    add_bullet_items(cell, [
        ("Historical tickets", "36 records \u2014 7 categories, 4 priorities, 3 channels"),
        ("KB articles", "12 articles with content, keywords, agent notes"),
        ("Escalation rules", "5 rules \u2192 5 teams (Eng, Integrations, Finance, Compliance, Platform)"),
        ("Eval set", "12 held-out cases (blind test \u2014 do NOT use for training)"),
        ("Pilot source", "Excel dataset"),
        ("Production source", "Freshworks suite APIs (post-pilot)"),
        ("PII", "None in pilot data; production requires handling policy"),
    ])

    # ============================================
    # ROW 2: Out of Scope (2 cols) | Feedback Loop & Guardrails (2 cols)
    # ============================================

    # --- Out of Scope ---
    cell = merge_cells_in_row(canvas, 2, 0, 1)
    set_cell_shading(cell, "F5F5F5")
    add_section_title(cell, "OUT OF SCOPE (PHASE 1)")
    add_bullet_items(cell, [
        "\u2716  Freshdesk API integration (architect for it, don't build)",
        "\u2716  Multi-language support (English only)",
        "\u2716  Customer-facing AI (agent-facing only)",
        "\u2716  Auto-send / auto-reply (human always approves)",
        "\u2716  Live KB refresh (static Excel data)",
        "\u2716  Production deployment / scaling",
    ])

    # --- Feedback Loop & Guardrails ---
    cell = merge_cells_in_row(canvas, 2, 2, 3)
    set_cell_shading(cell, "E8F0F8")
    add_section_title(cell, "FEEDBACK LOOP & GUARDRAILS", color=MEDIUM_BLUE)
    add_bullet_items(cell, [
        ("Feedback", "Agent rates + edits response \u2192 saved as correct answer \u2192 system learns"),
        ("Profanity", "Zero tolerance \u2014 all outputs checked"),
        ("Adversarial", "Graceful handling of misuse and edge cases"),
        ("Fallback", "'No confident answer' disclaimer when confidence is low"),
        ("Audit", "All copilot decisions logged for traceability"),
    ])

    # tight spacer
    spacer2 = doc.add_paragraph()
    spacer2.space_before = Pt(2)
    spacer2.space_after = Pt(0)

    # --- Footer ---
    footer_table = doc.add_table(rows=1, cols=3)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = footer_table.cell(0, 0)
    p = cell.paragraphs[0]
    add_formatted_text(p, "Gyde AI POD  |  ", bold=True, size=8, color=MEDIUM_GRAY)
    add_formatted_text(p, "Confidential", size=8, color=MEDIUM_GRAY)

    cell = footer_table.cell(0, 1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "Client: Prasanna  |  POD Lead: Amit  |  PM: Shivani", size=8, color=MEDIUM_GRAY)

    cell = footer_table.cell(0, 2)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_formatted_text(p, "Doc 03, Section 3.2", size=8, color=MEDIUM_GRAY)

    for cell in footer_table.rows[0].cells:
        set_cell_border(cell,
            top={"sz": "4", "color": "CCCCCC"},
            bottom={"sz": "0", "color": "FFFFFF"},
            left={"sz": "0", "color": "FFFFFF"},
            right={"sz": "0", "color": "FFFFFF"})

    return doc


if __name__ == "__main__":
    doc = build_canvas()

    # Save to client-docs folder
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "client-docs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Use_Case_Canvas_AI_Support_Copilot.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")

    # Also save to Simulation Docs
    sim_dir = "/Users/amit/Work/Gyde/Gyde Pivot/Gyde AI POD Framework/Simulation Docs"
    sim_path = os.path.join(sim_dir, "Use_Case_Canvas_AI_Support_Copilot.docx")
    doc.save(sim_path)
    print(f"Saved: {sim_path}")
